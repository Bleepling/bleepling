from __future__ import annotations
import tkinter as tk
from tkinter import ttk, filedialog
from pathlib import Path
import shutil, subprocess, json, sys, os, re
from difflib import SequenceMatcher

try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

try:
    from PIL import Image, ImageTk
except Exception:
    Image = None
    ImageTk = None
from bleepling.services.bleeping_service import BleepingService
from bleepling.services.time_service import format_time_point
from bleepling.utils.help_dialog import show_help_dialog


def _safe_read_lines(path: Path):
    if not path.exists():
        return []
    return [line.rstrip("\n") for line in path.read_text(encoding="utf-8", errors="ignore").splitlines()]


def _safe_write_lines(path: Path, lines):
    path.write_text("\n".join(lines).strip() + ("\n" if lines else ""), encoding="utf-8")


def _write_point_times_lines(path: Path, lines):
    _safe_write_lines(path, lines)


def _list_files(directory: Path, exts: set[str]):
    if not directory.exists():
        return []
    return sorted([p.name for p in directory.iterdir() if p.is_file() and p.suffix.lower() in exts], key=str.lower)


def _format_ts(start: float) -> str:
    return format_time_point(start, clamp_zero=False)


TITLES = {"Herr","Frau","Dr","Doktor","Professor","Prof","Richter","Richterin","Minister","Ministerin","Staatsminister","Präsident","Präsidentin"}
STOP = {
    "Und","Oder","Sie","Wir","Ich","Der","Die","Das","Ein","Eine","Den","Dem","Des","Nicht","Auch","Nur",
    "Tech","Deep","Dive","Google","Recherche","Theorie","Praxis","Token","Vektor","Temperatur","Worst","Case",
    "Montag","Dienstag","Mittwoch","Donnerstag","Freitag","ULG","RVG","StGB","KI","ChatGPT","Prompt","Prompts",
    "Experiment","Experimente","Grundlagen","Teil","Fusion","Frage","Fragen","Fragestellungen","Fällen","Fälle",
    "Bitte","Danke","Szenario","Szenarien","Geschichte","Regeln","Wahrscheinlichkeit","Verfahren","Videos"
}


def _clean_token(token: str) -> str:
    token = str(token).strip()
    token = re.sub(r"^[^\wÄÖÜäöüß]+|[^\wÄÖÜäöüß-]+$", "", token)
    return token


def _looks_like_name(token: str) -> bool:
    if not token or len(token) < 3:
        return False
    if token in STOP:
        return False
    if not token[:1].isupper():
        return False
    if token.isupper() and len(token) <= 4:
        return False
    return True


def _fuzzy_ratio(a: str, b: str) -> int:
    return int(round(100 * SequenceMatcher(None, a.lower(), b.lower()).ratio()))


def _best_match(value: str, items: list[str]):
    best_item = ""
    best_score = 0
    for item in items:
        score = _fuzzy_ratio(value, item)
        if score > best_score:
            best_item = item
            best_score = score
    return best_item, best_score





def _normalize_header_token(value: str) -> str:
    value = value.strip().lower()
    value = value.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    return re.sub(r"\s+", " ", value)

def _pdfplumber_tables_to_text(path: Path) -> str:
    texts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                texts.append(txt)
    return "\n".join(texts)

def _extract_names_from_pdf_tables(path: Path, include_firstnames: bool = False) -> list[str]:
    if pdfplumber is None:
        return []

    surnames: list[str] = []
    firstnames: list[str] = []

    def _clean_cell(value) -> str:
        if value is None:
            return ""
        return " ".join(str(value).replace("\n", " ").split()).strip()

    def _find_header_positions(row: list[str]) -> tuple[int | None, int | None]:
        first_idx = None
        last_idx = None
        for idx, cell in enumerate(row):
            norm = _normalize_header_token(cell)
            if "vorname" in norm and first_idx is None:
                first_idx = idx
            if "nachname" in norm and last_idx is None:
                last_idx = idx
        return first_idx, last_idx

    def _looks_like_number(value: str) -> bool:
        return bool(re.fullmatch(r"\d+[A-ZÄÖÜ]*", value.strip()))

    def _looks_like_land(value: str) -> bool:
        value = value.strip().upper()
        return bool(re.fullmatch(r"[A-ZÄÖÜ]{1,4}", value))

    def _is_plausible_person_name(value: str, allow_short_single_token: bool = False) -> bool:
        parts = [_clean_token(part) for part in value.split() if _clean_token(part)]
        if not parts:
            return False
        if any(_looks_like_name(part) for part in parts):
            return True
        if allow_short_single_token and len(parts) == 1:
            token = parts[0]
            if len(token) >= 2 and token[:1].isupper() and token.lower() not in {"bb", "be", "bw", "by", "hb", "he", "hh", "mv", "ni", "nw", "rp", "sh", "sl", "sn", "st", "th"}:
                return True
        return False

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []

            for table in tables:
                if not table:
                    continue

                header_positions = (None, None)
                data_started = False

                for raw_row in table:
                    row = [_clean_cell(cell) for cell in (raw_row or [])]
                    if not any(row):
                        continue

                    first_idx, last_idx = _find_header_positions(row)
                    if first_idx is not None and last_idx is not None:
                        header_positions = (first_idx, last_idx)
                        data_started = True
                        continue

                    if not data_started:
                        continue

                    first_idx, last_idx = header_positions
                    if first_idx is None or last_idx is None:
                        continue
                    if max(first_idx, last_idx) >= len(row):
                        continue

                    firstname = row[first_idx].strip()
                    surname = row[last_idx].strip()

                    # Manche PDFs verschieben laufende Nummer/Ländercode in Nachbarspalten.
                    # Für die Namensspalten selbst akzeptieren wir nur plausible Personennamen.
                    if not surname and last_idx + 1 < len(row):
                        candidate = row[last_idx + 1].strip()
                        if _is_plausible_person_name(candidate):
                            surname = candidate

                    if not firstname and first_idx - 1 >= 0:
                        candidate = row[first_idx - 1].strip()
                        if _is_plausible_person_name(candidate) and not _looks_like_land(candidate) and not _looks_like_number(candidate):
                            firstname = candidate

                    if not _is_plausible_person_name(surname, allow_short_single_token=True):
                        continue

                    surname_parts = [_clean_token(part) for part in surname.split() if _clean_token(part)]
                    if not surname_parts:
                        continue
                    surnames.append(" ".join(surname_parts))

                    if include_firstnames and _is_plausible_person_name(firstname):
                        first_parts = [_clean_token(part) for part in firstname.split() if _clean_token(part)]
                        if first_parts:
                            firstnames.append(" ".join(first_parts))

    return sorted(set(surnames + firstnames), key=str.lower)

def _extract_names_from_pdf_table(path: Path, include_firstnames: bool = False) -> list[str]:
    # Zuerst echte Tabellenzellen lesen. Das ist bei PDF-Formularen robuster als Fließtext,
    # weil Spalten wie Vorname/Nachname dort meist sauberer erhalten bleiben.
    table_names = _extract_names_from_pdf_tables(path, include_firstnames=include_firstnames)
    if table_names:
        return table_names

    raw_text = _extract_text_from_supported_file(path)
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]

    start_idx = None
    for i, ln in enumerate(lines):
        norm = _normalize_header_token(ln)
        if "vorname" in norm and "nachname" in norm:
            start_idx = i + 1
            break

    if start_idx is None:
        raise RuntimeError("Im PDF wurde keine Kopfzeile mit 'Vorname' und 'Nachname' gefunden.")

    names = []
    firstnames = []

    role_words = {
        "richter", "richterin", "staatsanwalt", "staatsanwältin", "oberstaatsanwältin",
        "vorsitzender", "vorsitzende", "direktorin", "direktor", "am", "beim", "des",
        "der", "im", "landgericht", "amtsgericht", "bundessozialgericht", "bundesgerichtshof",
        "verwaltungsgericht", "sozialgericht", "landessozialgericht", "lsg", "vg", "ag", "lg",
        "vors.", "probe", "oberregierungsrätin", "generalbundesanwalt", "bundesanwalt",
        "staatsanwaeltin", "oberstaatsanwaeltin", "richter", "richterin",
        "präsident", "praesident", "vizepräsident", "vizepraesident", "präsidentin", "praesidentin",
        "vizepräsidentin", "vizepraesidentin", "ministerialdirigent", "ministerialdirigentin",
        "leitender", "leitende", "oberrichter", "senatspräsident", "senatspraesident",
        "direktor", "direktorin", "vizedirektor", "vizedirektorin",
        "richter", "richterin", "vorsitzender", "vorsitzende", "vorsitzenderrichter", "vorsitzenderichterin",
        "rechtsanwalt", "rechtsanwältin", "rechtsanwaeltin", "notar", "notarin",
        "rechtspfleger", "rechtspflegerin", "oberstaatsanwalt", "oberstaatsanwältin", "oberstaatsanwaeltin",
        "staatsanwalt", "staatsanwältin", "staatsanwaeltin", "generalstaatsanwalt", "generalstaatsanwältin",
        "generalstaatsanwaeltin", "justiziar", "justiziarin", "dezernent", "dezernentin",
        "abteilungsleiter", "abteilungsleiterin", "referent", "referentin", "vors.", "präsidentdes", "richteram",
    }

    non_name_markers = {
        "teilnehmende", "damen", "herren", "diverse", "summe", "stand", "tn-anzahl",
        "links", "tagungsunterlagen", "gesamtanzahl", "veranstaltende", "ansprechpartner",
        "technische", "tagungsleitung", "thema", "zeitraum", "telefon", "telefax", "email",
        "e-mail-adresse", "private", "dienstliche", "bund", "bb", "be", "bw", "by", "hb",
        "he", "hh", "mv", "ni", "nw", "rp", "sh", "sl", "sn", "st", "th", "gast",
        "nürnberg", "nuernberg", "berlin", "potsdam", "cottbus", "frankfurt", "oder",
        "trier", "wustrau", "mainz", "koblenz", "zweibrücken", "zweibruecken",
    }

    def looks_like_data_line(ln: str) -> bool:
        # Allgemeiner als bisher:
        # Zeile beginnt mit laufender Nummer und enthält genügend Material für Teilnehmerdaten.
        if not re.match(r"^\d+\b", ln):
            return False
        tokens = [t for t in re.split(r"\s+", ln) if t]
        if len(tokens) < 4:
            return False
        return True

    def _surname_is_plausible(surname: str) -> bool:
        low = _normalize_header_token(surname)
        if low in role_words or low in non_name_markers:
            return False
        # einzelne Orts-/Funktionswörter abweisen
        blocked = {
            "präsident", "praesident", "vizepräsident", "vizepraesident", "präsidentin", "praesidentin",
            "vizepräsidentin", "vizepraesidentin", "direktor", "direktorin", "richter", "richterin",
            "rechtsanwalt", "rechtsanwältin", "rechtsanwaeltin", "notar", "notarin",
            "staatsanwalt", "staatsanwältin", "staatsanwaeltin", "oberstaatsanwalt", "oberstaatsanwältin",
            "oberstaatsanwaeltin", "rechtspfleger", "rechtspflegerin", "justiziar", "justiziarin",
            "nürnberg", "nuernberg", "berlin", "potsdam", "cottbus", "frankfurt", "oder",
            "trier", "wustrau", "mainz", "koblenz",
        }
        if low in blocked:
            return False
        # keine typischen Funktionsendungen als nacktes Wort
        if low.endswith(("gericht", "gerichtshof", "staatsanwaltschaft", "ministerium", "verwaltung")):
            return False
        return _looks_like_name(surname)

    def extract_from_line(ln: str):
        tokens = [t for t in re.split(r"\s+", ln) if t]
        if len(tokens) < 4:
            return None, None

        idx = 1  # nach laufender Nummer

        # optionales Land überspringen
        if idx < len(tokens) and re.fullmatch(r"[A-ZÄÖÜ]{1,4}", tokens[idx]):
            idx += 1

        # optionale Anrede
        if idx < len(tokens) and tokens[idx].lower() in {"herr", "frau"}:
            idx += 1

        # optionale Titel mehrfach
        while idx < len(tokens) and tokens[idx].rstrip(".").lower() in {"dr", "prof", "jr", "sr"}:
            idx += 1

        remaining = tokens[idx:]
        if len(remaining) < 2:
            return None, None

        name_tokens = []
        for tok in remaining:
            clean = _clean_token(tok)
            if not clean:
                continue
            low = _normalize_header_token(clean)

            if low in role_words:
                break
            if low in non_name_markers:
                break
            if re.fullmatch(r"\d+", clean):
                break

            # nur plausible Namenswörter
            if not re.search(r"[A-Za-zÄÖÜäöüß]", clean):
                break

            name_tokens.append(clean)
            if len(name_tokens) >= 5:
                break

        if len(name_tokens) < 2:
            return None, None

        firstname = " ".join(name_tokens[:-1]).strip()
        surname = name_tokens[-1].strip()

        if not _surname_is_plausible(surname):
            return None, None

        # Falls Vorname komplett unplausibel ist, nur Nachname übernehmen
        if firstname:
            first_parts = firstname.split()
            if not any(_looks_like_name(p) for p in first_parts):
                firstname = ""

        return firstname, surname

    for ln in lines[start_idx:]:
        if not looks_like_data_line(ln):
            continue
        firstname, surname = extract_from_line(ln)
        if surname:
            names.append(surname)
        if include_firstnames and firstname:
            firstnames.append(firstname)

    result = sorted(set(names + firstnames), key=str.lower)
    if not result:
        raise RuntimeError("Im PDF wurden nach der Kopfzeile keine verwertbaren Teilnehmer-Nachnamen gefunden.")
    return result
def _extract_text_from_supported_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".xlsx":
        if openpyxl is None:
            raise RuntimeError("Für XLSX-Dateien wird openpyxl benötigt.")
        wb = openpyxl.load_workbook(path, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if vals:
                    lines.append(" ".join(vals))
        return "\n".join(lines)
    if suffix == ".docx":
        if Document is None:
            raise RuntimeError("Für DOCX-Dateien wird python-docx benötigt.")
        doc = Document(str(path))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                vals = [c.text.strip() for c in row.cells if c.text.strip()]
                if vals:
                    lines.append(" ".join(vals))
        return "\n".join(lines)
    if suffix == ".pdf":
        if pdfplumber is not None:
            return _pdfplumber_tables_to_text(path)
        if PdfReader is None:
            raise RuntimeError("Für PDF-Dateien wird pdfplumber oder pypdf oder PyPDF2 benötigt.")
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                texts.append(txt)
        return "\n".join(texts)
    raise RuntimeError("Unterstützte Formate sind TXT, CSV, XLSX, DOCX und PDF.")


def _extract_name_tokens_from_text(raw_text: str, surnames_only: bool = True, include_firstnames: bool = False) -> list[str]:
    names = []
    for line in raw_text.splitlines():
        line = line.strip()
        if not line:
            continue
        parts = re.split(r"[;,|\t]+", line)
        for part in parts:
            tokens = [_clean_token(t) for t in re.split(r"\s+", part) if _clean_token(t)]
            cap_tokens = [t for t in tokens if _looks_like_name(t)]
            if not cap_tokens:
                continue

            # Standard: nur Nachnamen
            if surnames_only:
                last = cap_tokens[-1]
                names.append(last)
                if include_firstnames and len(cap_tokens) >= 2:
                    names.extend(cap_tokens[:-1])
            else:
                if include_firstnames:
                    names.extend(cap_tokens)
                else:
                    names.append(cap_tokens[-1])

    cleaned = []
    for n in names:
        n = n.strip()
        if not n:
            continue
        if len(n) < 2:
            continue
        cleaned.append(n)
    return sorted(set(cleaned), key=str.lower)

class BleepingTab(ttk.Frame):
    def __init__(self, master, app, set_status=None):
        super().__init__(master)
        self.app = app
        self._external_set_status = set_status
        self.progress_win = None
        self.progress_img = None
        self.bleeping_service = BleepingService()
        self._participant_import_surnames: list[str] = []
        self._participant_import_firstnames: list[str] = []
        self._participant_manual_baseline: list[str] = []
        self._participant_import_source_path: str = ""
        self._participant_option_trace_guard = False
        self._build()

    def _asset(self, name: str) -> Path:
        return Path(__file__).resolve().parents[3] / "assets" / name

    def _set_status(self, msg: str):
        if callable(self._external_set_status):
            self._external_set_status(msg)
        elif hasattr(self.app, "set_status"):
            self.app.set_status(msg)
        if hasattr(self, "status_box"):
            self.status_box.delete("1.0", "end")
            self.status_box.insert("1.0", msg)

    def _show_progress_window(self):
        if self.progress_win is not None:
            return
        win = tk.Toplevel(self)
        win.title("Bleepling arbeitet")
        win.transient(self.winfo_toplevel())
        try:
            win.attributes("-topmost", True)
        except Exception:
            pass
        win.resizable(False, False)
        frame = ttk.Frame(win, padding=18)
        frame.pack(fill="both", expand=True)

        bird_path = self._asset("vogel2_light_512_fixed.png")
        if Image is not None and ImageTk is not None and bird_path.exists():
            img = Image.open(bird_path).resize((180, 180))
            self.progress_img = ImageTk.PhotoImage(img)
            ttk.Label(frame, image=self.progress_img).pack(pady=(0, 10))
        else:
            ttk.Label(frame, text="🐤", font=("Segoe UI Emoji", 36)).pack(pady=(0, 10))

        msg = (
            "Bleepling bei der Arbeit - das könnte ein Weilchen dauern.\n"
            "Bitte warte, bis der Button nicht mehr rot ist, dann geht es weiter."
        )
        ttk.Label(frame, text=msg, justify="center").pack()

        win.update_idletasks()
        root = self.winfo_toplevel()
        try:
            rx = root.winfo_rootx()
            ry = root.winfo_rooty()
            rw = root.winfo_width()
            rh = root.winfo_height()
            ww = win.winfo_width()
            wh = win.winfo_height()
            x = rx + max(0, (rw - ww) // 2)
            y = ry + max(0, (rh - wh) // 2)
            win.geometry(f"+{x}+{y}")
        except Exception:
            pass

        self.progress_win = win
        try:
            self.app.set_running(True)
        except Exception:
            pass
        win.update()

    def _hide_progress_window(self):
        try:
            if self.progress_win is not None:
                self.progress_win.destroy()
        except Exception:
            pass
        self.progress_win = None
        try:
            self.app.set_running(False)
        except Exception:
            pass


    def _toggle_help(self, attr_name: str):
        help_text = getattr(self, attr_name, "")
        if not help_text:
            return
        show_help_dialog(self, "Hilfe", help_text)

    def _read_block_entries_from_widget(self) -> list[str]:
        return [x.strip() for x in self.block_text.get("1.0", "end").splitlines() if x.strip()]

    def _current_participant_import_entries(self) -> list[str]:
        selected: list[str] = []
        if self.participant_surnames_only.get():
            selected.extend(self._participant_import_surnames)
        if self.participant_include_firstnames.get():
            selected.extend(self._participant_import_firstnames)
        return sorted(set(selected), key=str.lower)

    def _extract_manual_block_entries(self) -> list[str]:
        if self._participant_manual_baseline:
            return list(self._participant_manual_baseline)
        current_entries = self._read_block_entries_from_widget()
        imported = {entry.casefold() for entry in (self._participant_import_surnames + self._participant_import_firstnames)}
        if not imported:
            return current_entries
        return [entry for entry in current_entries if entry.casefold() not in imported]

    def _persist_participant_import_state(self) -> None:
        if not self.app.project:
            return
        payload: dict[str, object] = {}
        if self.app.project.app_state_file.exists():
            try:
                payload = json.loads(self.app.project.app_state_file.read_text(encoding="utf-8") or "{}")
            except json.JSONDecodeError:
                payload = {}
        participant_state = {
            "surnames": list(self._participant_import_surnames),
            "firstnames": list(self._participant_import_firstnames),
            "manual_baseline": list(self._participant_manual_baseline),
            "source_path": self._participant_import_source_path,
            "include_firstnames": bool(self.participant_include_firstnames.get()),
        }
        payload["participant_import"] = participant_state
        self.app.project.app_state_file.parent.mkdir(parents=True, exist_ok=True)
        self.app.project.app_state_file.write_text(
            json.dumps(payload, indent=2, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )

    def _load_participant_import_state(self) -> None:
        self._participant_import_surnames = []
        self._participant_import_firstnames = []
        self._participant_manual_baseline = []
        if not self.app.project or not self.app.project.app_state_file.exists():
            return
        try:
            payload = json.loads(self.app.project.app_state_file.read_text(encoding="utf-8") or "{}")
        except json.JSONDecodeError:
            return
        participant_state = payload.get("participant_import", {})
        if not isinstance(participant_state, dict):
            return
        surnames = participant_state.get("surnames", [])
        firstnames = participant_state.get("firstnames", [])
        manual_baseline = participant_state.get("manual_baseline", [])
        source_path = participant_state.get("source_path", "")
        if isinstance(surnames, list):
            self._participant_import_surnames = [str(x).strip() for x in surnames if str(x).strip()]
        if isinstance(firstnames, list):
            self._participant_import_firstnames = [str(x).strip() for x in firstnames if str(x).strip()]
        if isinstance(manual_baseline, list):
            self._participant_manual_baseline = [str(x).strip() for x in manual_baseline if str(x).strip()]
        if isinstance(source_path, str):
            self._participant_import_source_path = source_path.strip()
        try:
            self._participant_option_trace_guard = True
            self.participant_include_firstnames.set(bool(participant_state.get("include_firstnames", False)))
        except Exception:
            pass
        finally:
            self._participant_option_trace_guard = False

    def _compute_participant_import_variants(self, file_path: Path) -> tuple[list[str], list[str]]:
        if file_path.suffix.lower() == ".pdf":
            surnames = _extract_names_from_pdf_table(file_path, include_firstnames=False)
            names_with_firstnames = _extract_names_from_pdf_table(file_path, include_firstnames=True)
        else:
            raw_text = _extract_text_from_supported_file(file_path)
            surnames = _extract_name_tokens_from_text(
                raw_text,
                surnames_only=True,
                include_firstnames=False,
            )
            names_with_firstnames = _extract_name_tokens_from_text(
                raw_text,
                surnames_only=bool(self.participant_surnames_only.get()),
                include_firstnames=True,
            )
        surname_set = {name.casefold() for name in surnames}
        firstnames = [name for name in names_with_firstnames if name.casefold() not in surname_set]
        return (
            sorted(set(surnames), key=str.lower),
            sorted(set(firstnames), key=str.lower),
        )

    def _rebuild_participant_import_from_source(self, announce: bool = False) -> bool:
        source = self._participant_import_source_path.strip()
        if not source:
            return False
        file_path = Path(source)
        if not file_path.exists():
            return False
        manual_entries = self._extract_manual_block_entries()
        surnames, firstnames = self._compute_participant_import_variants(file_path)
        self._participant_import_surnames = surnames
        self._participant_import_firstnames = firstnames
        self._apply_participant_import_entries(manual_entries=manual_entries, announce=announce)
        return True

    def refresh_participant_import(self) -> None:
        if not self.app.project:
            self._set_status("Bitte zuerst ein Projekt laden.")
            return
        if not self._participant_import_source_path.strip():
            self._set_status("Bitte zuerst eine Teilnehmerliste importieren.")
            return
        if not self._rebuild_participant_import_from_source(announce=False):
            self._set_status("Die zuletzt importierte Teilnehmerliste konnte nicht erneut geladen werden.")
            return

        mode_txt = []
        if self.participant_surnames_only.get():
            mode_txt.append("Nachnamen")
        if self.participant_include_firstnames.get():
            mode_txt.append("Vornamen")
        mode_label = " + ".join(mode_txt) if mode_txt else "keine Namensbestandteile"
        self._set_status(
            "Teilnehmerliste aktualisiert.\n"
            f"Aktive Einträge aus der Teilnehmerliste: {len(self._current_participant_import_entries())} ({mode_label}).\n"
            "Manuelle Blocklist-Einträge wurden beibehalten."
        )

    def _apply_participant_import_entries(self, manual_entries: list[str] | None = None, announce: bool = False) -> None:
        if not self.app.project:
            return
        if manual_entries is None:
            manual_entries = self._extract_manual_block_entries()
        self._participant_manual_baseline = list(manual_entries)
        merged = sorted(set(manual_entries + self._current_participant_import_entries()), key=str.lower)
        self.block_text.delete("1.0", "end")
        self.block_text.insert("1.0", "\n".join(merged))
        self.block_text.see("1.0")
        _safe_write_lines(self.app.project.blocklist_file, merged)
        self._persist_participant_import_state()
        if announce:
            count = len(self._current_participant_import_entries())
            mode_txt = []
            if self.participant_surnames_only.get():
                mode_txt.append("Nachnamen")
            if self.participant_include_firstnames.get():
                mode_txt.append("Vornamen")
            mode_label = " + ".join(mode_txt) if mode_txt else "keine Namensbestandteile"
            self._set_status(
                f"Teilnehmerlisten-Modus aktualisiert.\n"
                f"Aktive Einträge aus der Teilnehmerliste: {count} ({mode_label})."
            )

    def _on_participant_option_changed(self) -> None:
        if self._participant_option_trace_guard:
            return
        if not self.app.project:
            return
        if not (self._participant_import_surnames or self._participant_import_firstnames):
            return
        if self._rebuild_participant_import_from_source(announce=True):
            return
        self._apply_participant_import_entries(announce=True)

    def _on_participant_option_var_changed(self, *_args) -> None:
        self._schedule_participant_option_changed()

    def _schedule_participant_option_changed(self) -> None:
        if self._participant_option_trace_guard:
            return
        self._participant_option_trace_guard = True
        self.after_idle(self._run_participant_option_change)

    def _run_participant_option_change(self) -> None:
        try:
            self._participant_option_trace_guard = False
            self._on_participant_option_changed()
        finally:
            self._participant_option_trace_guard = False

    def _build(self):

        self.video_var = tk.StringVar()
        self.wav_var = tk.StringVar()
        self.json_var = tk.StringVar()
        self.candidate_var = tk.StringVar()
        self.block_thr = tk.IntVar(value=88)
        self.allow_thr = tk.IntVar(value=96)
        self.participant_surnames_only = tk.BooleanVar(value=True)
        self.participant_include_firstnames = tk.BooleanVar(value=False)
        self.participant_surnames_only.trace_add("write", self._on_participant_option_var_changed)
        self.participant_include_firstnames.trace_add("write", self._on_participant_option_var_changed)

        top = ttk.LabelFrame(self, text="1) Vorbereitung")
        top.pack(fill="x", padx=8, pady=8)
        top.grid_columnconfigure(1, weight=1)
        top.grid_columnconfigure(4, weight=1)

        ttk.Label(top, text="Ausgangsvideo").grid(row=0, column=0, sticky="w", padx=8, pady=4)
        self.video_combo = ttk.Combobox(top, textvariable=self.video_var, width=28, state="readonly")
        self.video_combo.grid(row=0, column=1, sticky="we", padx=4, pady=4)
        ttk.Button(top, text="WAV aus Video erzeugen", command=self.make_wav, style="Accent.TButton").grid(row=0, column=2, sticky="w", padx=6, pady=4)

        ttk.Label(top, text="words.json").grid(row=0, column=3, sticky="w", padx=(14, 4), pady=4)
        self.json_combo = ttk.Combobox(top, textvariable=self.json_var, width=24, state="readonly")
        self.json_combo.grid(row=0, column=4, sticky="we", padx=4, pady=4)
        ttk.Button(top, text="Kandidaten aus words.json erzeugen", command=self.make_candidates, style="Accent.TButton").grid(row=0, column=5, sticky="e", padx=8, pady=4)

        ttk.Label(top, text="WAV").grid(row=1, column=0, sticky="w", padx=8, pady=4)
        self.wav_combo = ttk.Combobox(top, textvariable=self.wav_var, width=28, state="readonly")
        self.wav_combo.grid(row=1, column=1, sticky="we", padx=4, pady=4)
        ttk.Button(top, text="words.json aus WAV erzeugen", command=self.make_words_json_stub, style="Accent.TButton").grid(row=1, column=2, sticky="w", padx=6, pady=4)

        prep_helpbar = ttk.Frame(top)
        prep_helpbar.grid(row=1, column=3, columnspan=3, sticky="e", padx=8, pady=(2, 4))
        ttk.Label(prep_helpbar, text="Hilfe zu Vorbereitung").pack(side="left")
        ttk.Button(prep_helpbar, text="?", width=3, command=lambda: self._toggle_help("prep_help_label")).pack(side="left", padx=(6, 0))
        self.prep_help_label = "Ablauf in kurz: 1) Video oder WAV vorbereiten  2) words.json erzeugen  3) Kandidaten-Datei erzeugen  4) Kandidaten auswerten  5) In der Vorschau entscheiden  6) Times-Datei erzeugen.\n\nKandidaten-Datei = Trefferdatei mit Zeitstempel, Nachname und Kontext. Das ist der normale Arbeitsstand für die Prüfung.\n\nTeilnehmerliste = reine Namensliste ohne Zeitpunkte. Sie dient nur als Hilfe für die Blocklist-Vorlage und erzeugt noch keine Times-Datei.\n\nBei Teilnehmerlisten können Nachnamen, Vornamen oder beide gemeinsam übernommen werden."

        mid = ttk.Frame(self)
        mid.pack(fill="x", padx=8, pady=(0, 8))
        mid.grid_columnconfigure(0, weight=2)
        mid.grid_columnconfigure(1, weight=1)

        left = ttk.LabelFrame(mid, text="2) Namenslisten und Regeln")
        left.grid(row=0, column=0, sticky="nsew")
        right = ttk.LabelFrame(mid, text="3) Kandidaten prüfen")
        right.grid(row=0, column=1, sticky="nsew", padx=(8, 0))

        ll = ttk.Frame(left)
        ll.pack(fill="both", expand=True, padx=8, pady=8)
        ll.grid_columnconfigure(0, weight=1)
        ll.grid_columnconfigure(1, weight=1)

        blockf = ttk.Frame(ll)
        blockf.grid(row=0, column=0, sticky="nsew")
        allowf = ttk.Frame(ll)
        allowf.grid(row=0, column=1, sticky="nsew", padx=(8, 0))

        ttk.Label(blockf, text="Blocklist (optional)").pack(anchor="w")
        self.block_text = tk.Text(blockf, height=5, width=38, wrap="none", undo=True)
        self.block_text.pack(fill="both", expand=True)

        ttk.Label(allowf, text="Allowlist (optional)").pack(anchor="w")
        self.allow_text = tk.Text(allowf, height=5, width=38, wrap="none", undo=True)
        self.allow_text.pack(fill="both", expand=True)

        bar = ttk.Frame(left)
        bar.pack(fill="x", padx=8, pady=(0, 6))
        ttk.Button(bar, text="Blocklist aus Kandidaten-Datei füllen", command=self.fill_blocklist_from_candidates, style="Accent.TButton").pack(side="left")
        ttk.Button(bar, text="Projekt speichern", command=self.save_lists, style="Accent.TButton").pack(side="left", padx=8)

        participant_bar = ttk.Frame(left)
        participant_bar.pack(fill="x", padx=8, pady=(0, 6))
        ttk.Button(participant_bar, text="Teilnehmerliste importieren", command=self.import_participant_list, style="Accent.TButton").pack(side="left")
        self.participant_surnames_check = ttk.Checkbutton(participant_bar, text="Nachnamen", variable=self.participant_surnames_only, command=self._schedule_participant_option_changed)
        self.participant_surnames_check.pack(side="left", padx=(12, 8))
        self.participant_firstnames_check = ttk.Checkbutton(participant_bar, text="Vornamen", variable=self.participant_include_firstnames, command=self._schedule_participant_option_changed)
        self.participant_firstnames_check.pack(side="left")
        ttk.Button(participant_bar, text="Akt.", width=5, command=self.refresh_participant_import, style="Accent.TButton").pack(side="left", padx=(8, 0))
        for widget in (self.participant_surnames_check, self.participant_firstnames_check):
            widget.bind("<ButtonRelease-1>", lambda _e: self._schedule_participant_option_changed(), add="+")
            widget.bind("<KeyRelease-space>", lambda _e: self._schedule_participant_option_changed(), add="+")

        ttk.Label(participant_bar, text="").pack(side="left", expand=True, fill="x")
        ttk.Label(participant_bar, text="Hilfe zu Namenslisten und Regeln").pack(side="left", padx=(18, 6))
        ttk.Button(participant_bar, text="?", width=3, command=lambda: self._toggle_help("lists_help_label")).pack(side="left")

        self.lists_help_label = "Teilnehmerlisten können als TXT, CSV, XLSX, DOCX oder PDF importiert werden. Ergebnis ist immer nur eine Namensbasis für die Blocklist-Vorlage.\n\nDie Allowlist ist stärker: Ein dort eingetragener Name gewinnt bei der späteren Auswertung gegen die Blocklist.\n\nBlocklist = Namen, die besonders genau geprüft werden sollen. Allowlist = Namen oder Schreibweisen, die nicht gebleept werden sollen.\n\nBei Teilnehmerlisten können Nachnamen, Vornamen oder beide gemeinsam übernommen werden. Die Blocklist erzwingt keinen Bleep automatisch; maßgeblich ist immer die spätere Auswertung."

        for c in range(4):
            right.grid_columnconfigure(c, weight=0)
        right.grid_columnconfigure(1, weight=1)
        right.grid_columnconfigure(3, weight=1)
        right.grid_rowconfigure(3, weight=1)

        ttk.Label(right, text="Kandidaten-Datei").grid(row=0, column=0, sticky="w", padx=8, pady=(8, 4))
        self.candidate_combo = ttk.Combobox(right, textvariable=self.candidate_var, width=12, state="readonly")
        self.candidate_combo.grid(row=0, column=1, sticky="we", padx=4, pady=(8, 4))
        ttk.Button(right, text="Datei laden", command=self.choose_candidate, style="Accent.TButton").grid(row=0, column=2, padx=4, pady=(8, 4))
        ttk.Button(right, text="Liste aktualisieren", command=self.refresh, style="Accent.TButton").grid(row=0, column=3, padx=4, pady=(8, 4), sticky="e")

        ttk.Label(right, text="Blocklist-Fuzzy %").grid(row=1, column=0, sticky="w", padx=8, pady=4)
        ttk.Spinbox(right, from_=50, to=100, textvariable=self.block_thr, width=5).grid(row=1, column=1, sticky="w", padx=4)
        ttk.Label(right, text="Allowlist-Fuzzy %").grid(row=1, column=2, sticky="w", padx=8, pady=4)
        ttk.Spinbox(right, from_=50, to=100, textvariable=self.allow_thr, width=5).grid(row=1, column=3, sticky="w", padx=4)

        ttk.Button(right, text="Auswertung starten", command=self.evaluate, style="Accent.TButton").grid(row=2, column=0, columnspan=2, padx=8, pady=8, sticky="we")
        ttk.Button(right, text="Schnell-Nachbleepen", command=self.quick, style="Accent.TButton").grid(row=2, column=2, columnspan=2, padx=8, pady=8, sticky="we")

        candidate_helpbar = ttk.Frame(right)
        candidate_helpbar.grid(row=3, column=0, columnspan=4, sticky="se", padx=8, pady=(10, 8))
        ttk.Label(candidate_helpbar, text="Hilfe zu Kandidaten prüfen").pack(side="left")
        ttk.Button(candidate_helpbar, text="?", width=3, command=lambda: self._toggle_help("candidate_help_label")).pack(side="left", padx=(6, 0))

        self.candidate_help_label = "Auswertung starten vergleicht die Trefferdatei mit Blocklist und Allowlist und füllt damit die Vorschau.\n\nSchnell-Nachbleepen erzeugt nur zusätzliche Zeitpunkte für einen zweiten Korrekturdurchgang."

        prevf = ttk.LabelFrame(self, text="4) Vorschau, Entscheidung und Times-Datei")
        prevf.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        preview_bar = ttk.Frame(prevf)
        preview_bar.pack(fill="x", padx=8, pady=(8, 6))
        ttk.Button(preview_bar, text="Markierte löschen", command=self.delete_selected_rows, style="Accent.TButton").pack(side="left")
        ttk.Button(preview_bar, text="Markierte in Allowlist", command=self.add_selected_to_allowlist, style="Accent.TButton").pack(side="left", padx=8)
        ttk.Button(preview_bar, text="Markierte erlauben", command=self.mark_selected_allow, style="Accent.TButton").pack(side="left", padx=8)
        ttk.Button(preview_bar, text="Markierte bleepen", command=self.mark_selected_bleep, style="Accent.TButton").pack(side="left", padx=8)
        ttk.Button(preview_bar, text="Times-Datei aus Vorschau", command=self.create_times_from_preview, style="Accent.TButton").pack(side="left", padx=8)
        ttk.Label(preview_bar, text="").pack(side="left", expand=True, fill="x")
        ttk.Label(preview_bar, text="Hilfe zu Vorschau, Entscheidung und Times-Datei").pack(side="left", padx=(12, 6))
        ttk.Button(preview_bar, text="?", width=3, command=lambda: self._toggle_help("preview_help_label")).pack(side="left", padx=(0, 0))

        self.preview_help_label = "Hier wird entschieden, was am Ende tatsächlich gebleept wird. Erst mit \"Times-Datei aus Vorschau\" entsteht die spätere Bleepliste.\n\nIn Allowlist = Schreibweise merken. Erlauben = Treffer freigeben und aus der offenen Prüfung nehmen. Bleepen = Treffer für die Times-Datei vormerken."

        self.preview = ttk.Treeview(prevf, columns=("zeit", "kandidat", "ergebnis", "regel", "kontext"), show="headings", height=10, selectmode="extended")
        for c, w in [("zeit", 90), ("kandidat", 160), ("ergebnis", 90), ("regel", 220), ("kontext", 820)]:
            self.preview.heading(c, text=c.capitalize())
            self.preview.column(c, width=w, anchor="w")
        self.preview.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        self._preview_anchor = None
        self.preview.bind("<Button-1>", self._on_preview_click, add="+")
        self.preview.bind("<Control-a>", self._select_all_preview_rows, add="+")
        self.preview.bind("<Control-A>", self._select_all_preview_rows, add="+")
        try:
            self.bind_all("<Control-a>", self._select_all_preview_rows, add="+")
            self.bind_all("<Control-A>", self._select_all_preview_rows, add="+")
        except Exception:
            pass

        statusf = ttk.LabelFrame(self, text="Status und Hinweise")
        statusf.pack(fill="x", padx=8, pady=(0, 8))
        self.status_box = tk.Text(statusf, height=3, wrap="word")
        self.status_box.pack(fill="x", padx=8, pady=8)
        self.refresh()

    def refresh(self):
        if not self.app.project:
            return
        p = self.app.project
        self._load_participant_import_state()
        self.video_combo["values"] = _list_files(p.input_video_dir, {".mp4", ".mov", ".mkv", ".avi", ".m4v", ".wmv"})
        self.wav_combo["values"] = _list_files(p.transcription_wav_dir, {".wav"})
        self.json_combo["values"] = _list_files(p.transcription_json_dir, {".json"})
        self.candidate_combo["values"] = _list_files(p.candidates_raw_dir, {".txt"})
        if self.video_combo["values"] and (not self.video_var.get() or self.video_var.get() not in self.video_combo["values"]):
            self.video_var.set(self.video_combo["values"][0])
        if self.wav_combo["values"] and (not self.wav_var.get() or self.wav_var.get() not in self.wav_combo["values"]):
            self.wav_var.set(self.wav_combo["values"][0])
        if self.json_combo["values"] and (not self.json_var.get() or self.json_var.get() not in self.json_combo["values"]):
            self.json_var.set(self.json_combo["values"][0])
        if self.candidate_combo["values"] and (not self.candidate_var.get() or self.candidate_var.get() not in self.candidate_combo["values"]):
            self.candidate_var.set(self.candidate_combo["values"][0])
        if self._participant_import_surnames or self._participant_import_firstnames:
            displayed = sorted(set(self._participant_manual_baseline + self._current_participant_import_entries()), key=str.lower)
        else:
            displayed = _safe_read_lines(p.blocklist_file)
        self.block_text.delete("1.0", "end"); self.block_text.insert("1.0", "\n".join(displayed))
        self.allow_text.delete("1.0", "end"); self.allow_text.insert("1.0", "\n".join(_safe_read_lines(p.allowlist_file)))

    def save_lists(self):
        if not self.app.project:
            return
        p = self.app.project
        current_entries = [x.strip() for x in self.block_text.get("1.0", "end").splitlines() if x.strip()]
        if self._participant_import_surnames or self._participant_import_firstnames:
            imported = {entry.casefold() for entry in self._current_participant_import_entries()}
            self._participant_manual_baseline = [entry for entry in current_entries if entry.casefold() not in imported]
            self._persist_participant_import_state()
        _safe_write_lines(p.blocklist_file, current_entries)
        _safe_write_lines(p.allowlist_file, [x.strip() for x in self.allow_text.get("1.0", "end").splitlines() if x.strip()])
        self._set_status("Projekt gespeichert: Listen wurden gesichert.")


    def import_participant_list(self):
        if not self.app.project:
            self._set_status("Bitte zuerst ein Projekt laden.")
            return
        path = filedialog.askopenfilename(
            title="Teilnehmerliste auswählen",
            filetypes=[
                ("Unterstützte Dateien", "*.txt *.csv *.xlsx *.docx *.pdf"),
                ("Alle Dateien", "*.*"),
            ],
        )
        if not path:
            return
        try:
            file_path = Path(path)
            surnames, firstnames = self._compute_participant_import_variants(file_path)
            if not surnames:
                self._set_status("Es konnten aus der Teilnehmerliste keine plausiblen Namen gelesen werden.")
                return

            manual_entries = self._extract_manual_block_entries()
            self._participant_import_source_path = str(file_path)
            self._participant_import_surnames = sorted(set(surnames), key=str.lower)
            self._participant_import_firstnames = sorted(set(firstnames), key=str.lower)
            self._apply_participant_import_entries(manual_entries=manual_entries, announce=False)

            mode_txt = []
            if self.participant_surnames_only.get():
                mode_txt.append("Nachnamen")
            if self.participant_include_firstnames.get():
                mode_txt.append("Vornamen")
            mode_label = " + ".join(mode_txt) if mode_txt else "keine Namensbestandteile"
            self._set_status(
                f"Teilnehmerliste importiert: {Path(path).name}\n"
                f"Übernommen in Blocklist-Vorlage: {len(self._current_participant_import_entries())} Einträge ({mode_label}).\n"
                f"Hinweis: Das ist noch keine Kandidaten-Datei mit Zeitpunkten, sondern nur eine Namensbasis für die spätere Auswertung."
            )
        except Exception as e:
            self._set_status(f"Teilnehmerliste konnte nicht importiert werden: {e}")

    def choose_candidate(self):
        if not self.app.project:
            return
        path = filedialog.askopenfilename(filetypes=[("Textdateien", "*.txt")])
        if not path:
            return
        target = self.app.project.candidates_raw_dir / Path(path).name
        if Path(path) != target:
            shutil.copy2(path, target)
        self.refresh()
        self.candidate_var.set(target.name)
        self._set_status(f"Kandidaten-Datei gewählt: {target.name}")

    def make_wav(self):
        if not self.app.project or not self.video_var.get():
            self._set_status("Bitte zuerst Projekt und Videodatei wählen.")
            return
        video = self.app.project.input_video_dir / self.video_var.get()
        wav = self.app.project.transcription_wav_dir / (video.stem + ".wav")
        ffmpeg = shutil.which("ffmpeg") or shutil.which("ffmpeg.exe")
        if ffmpeg:
            try:
                subprocess.run([ffmpeg, "-y", "-i", str(video), "-vn", "-ac", "1", "-ar", "16000", str(wav)], check=True, capture_output=True)
                self.refresh()
                self.wav_var.set(wav.name)
                self._set_status(f"WAV erzeugt: {wav.name}")
                return
            except Exception as e:
                self._set_status(f"WAV-Erzeugung fehlgeschlagen: {e}")
                return
        self._set_status("FFmpeg wurde nicht gefunden.")

    def make_words_json_stub(self):
        if not self.app.project or not self.wav_var.get():
            self._set_status("Bitte zuerst eine WAV-Datei wählen.")
            return

        wav = self.app.project.transcription_wav_dir / self.wav_var.get()
        out = self.app.project.transcription_json_dir / (wav.stem + ".words.json")
        script = Path(__file__).resolve().parents[3] / "legacy_reference" / "v5_scripts" / "transcribe_with_word_timestamps_cuda_fix2.py"

        self._show_progress_window()
        self._set_status("Transkription gestartet ...")
        try:
            self.app.update_idletasks()
        except Exception:
            pass

        try:
            if script.exists():
                settings = {}
                try:
                    settings = self.app.project.read_settings()
                except Exception:
                    settings = {}

                model = settings.get("whisper_model", "medium")
                mode = settings.get("transcription_mode", "auto")
                device = "cuda" if mode in ("gpu", "auto") else "cpu"

                env = os.environ.copy()
                extra_paths = settings.get("extra_cuda_paths", "")
                if extra_paths:
                    env["PATH"] = extra_paths + ";" + env.get("PATH", "")

                cmd = [
                    sys.executable,
                    str(script),
                    "--input", str(wav),
                    "--output-dir", str(self.app.project.transcription_json_dir),
                    "--model", str(model),
                    "--device", device,
                    "--language", "de",
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, env=env)
                if result.returncode == 0 and out.exists():
                    self.refresh()
                    self.json_var.set(out.name)
                    self.make_candidates(auto_only=True)
                    self._set_status(f"words.json erzeugt: {out.name}")
                    return

            if not out.exists():
                out.write_text(json.dumps({"source_file": wav.name, "segments": []}, indent=2, ensure_ascii=False), encoding="utf-8")
            self.refresh()
            self.json_var.set(out.name)
            self.make_candidates(auto_only=True)
            self._set_status(f"Platzhalter-words.json erzeugt: {out.name}")
        except Exception as e:
            self._set_status(f"Transkription fehlgeschlagen: {e}")
        finally:
            self._hide_progress_window()

    def make_candidates(self, auto_only=False):
        try:
            if not self.app.project or not self.json_var.get():
                self._set_status("Bitte zuerst eine words.json wählen.")
                return

            json_path = self.app.project.transcription_json_dir / self.json_var.get()
            if not json_path.exists():
                self._set_status(f"words.json nicht gefunden: {json_path.name}")
                return

            out_path = self.bleeping_service.generate_candidate_file_from_words_json(self.app.project, json_path)
            self.refresh()
            self.candidate_var.set(out_path.name)

            if not auto_only:
                entry_count = len(self.bleeping_service.parse_candidate_file(out_path))
                self._set_status(f"Kandidaten-Datei erzeugt: {out_path.name}\nTreffer: {entry_count}")
        except Exception as e:
            self._set_status(f"Kandidaten aus words.json erzeugen fehlgeschlagen: {e}")

    def fill_blocklist_from_candidates(self):
        if not self.app.project or not self.candidate_var.get():
            self._set_status("Bitte zuerst eine Kandidaten-Datei wählen.")
            return
        path = self.app.project.candidates_raw_dir / self.candidate_var.get()
        names = []
        for line in _safe_read_lines(path):
            parts = [p.strip() for p in line.split("|", 2)]
            if len(parts) >= 3 and parts[1]:
                names.append(parts[1])
        names = sorted(set(names))
        self.block_text.delete("1.0", "end")
        self.block_text.insert("1.0", "\n".join(names[:200]))
        self._set_status(f"Blocklist aus Kandidaten-Datei gefüllt. Übernommen: {min(len(names), 200)} Einträge.")

    def evaluate(self):
        if not self.app.project or not self.candidate_var.get():
            self._set_status("Bitte zuerst eine Kandidaten-Datei wählen.")
            return

        path = self.app.project.candidates_raw_dir / self.candidate_var.get()
        blocklist_text = self.block_text.get("1.0", "end")
        allowlist_text = self.allow_text.get("1.0", "end")
        block_thr = int(self.block_thr.get())
        allow_thr = int(self.allow_thr.get())

        entries = self.bleeping_service.parse_candidate_file(path)
        decisions = self.bleeping_service.evaluate_candidates(
            entries=entries,
            blocklist_text=blocklist_text,
            allowlist_text=allowlist_text,
            blocklist_threshold=block_thr,
            allowlist_threshold=allow_thr,
        )

        for i in self.preview.get_children():
            self.preview.delete(i)

        visible_count = 0
        for item in decisions:
            if item.decision == "ignorieren":
                continue
            self.preview.insert(
                "",
                "end",
                values=(item.timestamp, item.candidate, item.decision, item.reason, item.context),
            )
            visible_count += 1

        summary = self.bleeping_service.summarize_decisions(decisions)
        self._set_status(
            f"Auswertung abgeschlossen. Vorschau: {visible_count} | Gesamt: {summary.get('gesamt', 0)} | "
            f"Bleepen: {summary.get('bleepen', 0)} | Erlaubt: {summary.get('erlaubt', 0)} | "
            f"Prüfen: {summary.get('prüfen', 0)} | Ignorieren: {summary.get('ignorieren', 0)}"
        )
    def _select_all_preview_rows(self, event=None):
        items = self.preview.get_children("")
        if items:
            self.preview.selection_set(items)
            self._preview_anchor = items[0]
        return "break"

    def _on_preview_click(self, event):
        region = self.preview.identify("region", event.x, event.y)
        if region not in {"tree", "cell"}:
            return

        iid = self.preview.identify_row(event.y)
        if not iid:
            return

        ctrl_pressed = bool(event.state & 0x0004)
        shift_pressed = bool(event.state & 0x0001)
        items = list(self.preview.get_children(""))

        if shift_pressed and self._preview_anchor in items:
            start = items.index(self._preview_anchor)
            end = items.index(iid)
            lo, hi = sorted((start, end))
            self.preview.selection_set(items[lo:hi + 1])
        elif ctrl_pressed:
            current = set(self.preview.selection())
            if iid in current:
                current.remove(iid)
            else:
                current.add(iid)
                self._preview_anchor = iid
            self.preview.selection_set(list(current))
        else:
            self.preview.selection_set((iid,))
            self._preview_anchor = iid

        self.preview.focus(iid)
        return "break"

    def _set_selected_result(self, result_text: str, rule_text: str):
        selected = self.preview.selection()
        if not selected:
            self._set_status("Keine Treffer markiert.")
            return 0, []
        count = 0
        selected_ids = list(selected)
        for iid in selected_ids:
            vals = list(self.preview.item(iid, "values"))
            if len(vals) >= 5:
                vals[2] = result_text
                vals[3] = rule_text
                self.preview.item(iid, values=vals)
                count += 1
        return count, selected_ids

    def mark_selected_bleep(self):
        count, _ = self._set_selected_result("bleepen", "Manuell bleepen gesetzt")
        self._set_status(f"Markierte Treffer bleepen: {count}")

    def mark_selected_allow(self):
        count, selected_ids = self._set_selected_result("erlaubt", "Manuell erlaubt und aus offener Liste entfernt")
        for iid in selected_ids:
            self.preview.delete(iid)
        self._set_status(f"Markierte Treffer erlaubt und entfernt: {count}")

    def add_selected_to_allowlist(self):
        selected = self.preview.selection()
        if not selected:
            self._set_status("Keine Treffer markiert.")
            return
        existing = [x.strip() for x in self.allow_text.get("1.0", "end").splitlines() if x.strip()]
        added = 0
        for iid in selected:
            vals = self.preview.item(iid, "values")
            if len(vals) >= 2:
                cand = str(vals[1]).strip()
                if cand and cand not in existing:
                    existing.append(cand)
                    added += 1
        self.allow_text.delete("1.0", "end")
        self.allow_text.insert("1.0", "\n".join(existing))
        self._set_status(f"Markierte Treffer in Allowlist übernommen: {added}")

    def delete_selected_rows(self):
        selected = list(self.preview.selection())
        for iid in selected:
            self.preview.delete(iid)
        self._set_status(f"Markierte Treffer gelöscht: {len(selected)}")

    def _current_preview_rows(self):
        rows = []
        for iid in self.preview.get_children():
            vals = self.preview.item(iid, "values")
            if vals and len(vals) >= 5:
                rows.append(tuple(vals))
        return rows

    def create_point_times_from_preview(self):
        if not self.app.project or not self.candidate_var.get():
            self._set_status("Bitte zuerst eine Kandidaten-Datei wählen.")
            return

        stem = Path(self.candidate_var.get()).stem
        times_path = self.app.project.times_dir / f"{stem}.times.txt"
        reviewed_path = self.app.project.candidates_reviewed_dir / f"{stem}.reviewed.txt"

        rows = self._current_preview_rows()
        times_lines = []
        reviewed_lines = []

        for ts, cand, result, rule, ctx in rows:
            reviewed_lines.append(f"{ts} | {cand} | {result} | {rule} | {ctx}")
            if str(result).strip().lower() == "bleepen":
                times_lines.append(str(ts).strip())

        _write_point_times_lines(times_path, times_lines)
        _safe_write_lines(reviewed_path, reviewed_lines)
        self._set_status(
            f"Times-Datei erzeugt: {times_path.name}\n"
            f"Bleeps: {len(times_lines)} | Reviewed-Zeilen: {len(reviewed_lines)}"
        )

    def create_times_from_preview(self):
        # Transitional compatibility wrapper: this preview path still writes
        # point-based .times.txt files for the existing workflow.
        self.create_point_times_from_preview()

    def quick(self):
        if not self.app.project or not self.candidate_var.get():
            self._set_status("Bitte zuerst eine Kandidaten-Datei wählen.")
            return

        stem = Path(self.candidate_var.get()).stem
        times_path = self.app.project.times_dir / f"{stem}.times.txt"
        quick_path = self.app.project.times_dir / f"{stem}.quick.times.txt"

        current = [x.strip() for x in _safe_read_lines(times_path) if x.strip()]
        existing_quick = [x.strip() for x in _safe_read_lines(quick_path) if x.strip()]
        existing_set = set(existing_quick)
        new_items = [x for x in current if x not in existing_set]
        merged = existing_quick + new_items
        _safe_write_lines(quick_path, merged)

        self._set_status(
            f"Schnell-Nachbleepen vorbereitet: {quick_path.name}\n"
            f"Neue Zeitpunkte: {len(new_items)}"
        )
